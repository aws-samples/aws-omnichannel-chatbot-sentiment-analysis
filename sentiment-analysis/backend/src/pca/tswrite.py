"""
Produce Word Document transcriptions using the automatic speech recognition from AWS Transcribe.
Requires the following non-standard Amazon Linux libraries:
"""

from parsets import TranscribeParser
from docx import Document
from docx.table import Table
from docx.shared import Cm, Mm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from pathlib import Path
from time import perf_counter
from scipy.interpolate import make_interp_spline
import json, datetime
import matplotlib.pyplot as plt
import numpy as np
import statistics
import sys
import parsets
import os

# Column offsets in Transcribe output document table
COL_STARTTIME = 0
COL_ENDTIME = 1
COL_SPEAKER = 2
COL_CONTENT = 3
COL_SENTIMENT = 4
COL_SENTIMENT_SCORE = 5

# Sentiment helpers
MIN_SENTIMENT_NEGATIVE = 0.4
MIN_SENTIMENT_POSITIVE = 0.4
COMPREHEND_ENTITY = "onsei"

def convertTimeStamp(timeInSeconds: float) -> str:
    """ Function to help convert timestamps from s to H:M:S:MM """
    timeDelta = datetime.timedelta(seconds=float(timeInSeconds))
    tsFront = timeDelta - datetime.timedelta(microseconds=timeDelta.microseconds)
    tsSmall = timeDelta.microseconds
    return str(tsFront) + "." + str(int(tsSmall / 10000))


def getTextColour(confidence: float) -> RGBColor:
    """
    Get an RGB colour to represent this confidence range - it's a float,
    so can't easily be a simple lookup.  This is used for any work rendering
    """
    if confidence >= 0.95:
        textColour = RGBColor(0, 0, 0)
    elif confidence >= 0.90:
        textColour = RGBColor(32, 0, 0)
    elif confidence >= 0.75:
        textColour = RGBColor(64, 0, 0)
    elif confidence >= 0.6:
        textColour = RGBColor(128, 0, 0)
    elif confidence >= 0.5:
        textColour = RGBColor(191, 0, 0)
    else:
        textColour = RGBColor(255, 0, 0)

    return textColour


def addNextConfidenceRow(confTable: Table, rowName: str, rowStatsLookup: int, parsedWords: int):
    # Repeated code to build up the word confidence table
    row_cells = confTable.add_row().cells
    row_cells[0].text = rowName
    row_cells[1].text = str(rowStatsLookup)
    row_cells[2].text = str(round(rowStatsLookup / parsedWords * 100, 2)) + "%"


def writeOutTranscribeTable(outputTable, transcribeSegments):
    # Create a row populate it for each segment that we have
    for segment in transcribeSegments:

        # Start with the easy stuff
        row_cells = outputTable.add_row().cells
        row_cells[COL_STARTTIME].text = convertTimeStamp(segment["SegmentStartTime"])
        row_cells[COL_ENDTIME].text = convertTimeStamp(segment["SegmentEndTime"])
        row_cells[COL_SPEAKER].text = segment["SegmentSpeaker"]

        # Then do each word with confidence-level colouring
        for eachWord in segment["WordConfidence"]:
            run = row_cells[COL_CONTENT].paragraphs[0].add_run(eachWord["Text"])
            confLevel = eachWord["Confidence"]
            run.font.color.rgb = getTextColour(confLevel)
            # If the confidence is < 75% then additionally highlight it
            if confLevel < 0.75:
                run.font.highlight_color = WD_COLOR_INDEX.GRAY_25

        # Finish with the base sentiment for the segment - don't write out score
        # if it turns out that this segment ie neither Negative nor Positive
        if bool(segment["SentimentIsPositive"]):
            row_cells[COL_SENTIMENT].text = "Positive"
            row_cells[COL_SENTIMENT_SCORE].text = str(segment["SentimentScore"])[:4]
        elif bool(segment["SentimentIsNegative"]):
            row_cells[COL_SENTIMENT].text = "Negative"
            row_cells[COL_SENTIMENT_SCORE].text = str(segment["SentimentScore"])[:4]


def write(inputFilename, docxFilename, transcribeParser):
    """
    Write a transcript from the .json transcription file and other data generated
    by the results parser, putting it all into a human-readable Word document
    """
    json_filepath = Path(inputFilename)
    parseJobInfo = json.load(open(json_filepath.absolute(), "r", encoding="utf-8"))
    analysisJobInfo = parseJobInfo["ConversationAnalytics"]
    speechSegmentList = parseJobInfo["SpeechSegments"]
    transcribeJjobInfo = analysisJobInfo["SourceInformation"][0]["TranscribeJobInfo"]

    # Stats dictionary
    stats = {
        "timestamps": [],
        "accuracy": [],
        "9.8": 0, "9": 0, "8": 0, "7": 0, "6": 0, "5": 0, "4": 0, "3": 0, "2": 0, "1": 0, "0": 0,
        "parsedWords": 0 }

    # Word accuracy count
    for segment in speechSegmentList:
        for confidenceList in segment["WordConfidence"]:
            stats["timestamps"].append(float(confidenceList["StartTime"]))
            wordConf = float(confidenceList["Confidence"])
            stats["accuracy"].append(wordConf * 100.0)
            if wordConf >= 0.98: stats["9.8"] += 1
            elif wordConf >= 0.9: stats["9"] += 1
            elif wordConf >= 0.8: stats["8"] += 1
            elif wordConf >= 0.7: stats["7"] += 1
            elif wordConf >= 0.6: stats["6"] += 1
            elif wordConf >= 0.5: stats["5"] += 1
            elif wordConf >= 0.4: stats["4"] += 1
            elif wordConf >= 0.3: stats["3"] += 1
            elif wordConf >= 0.2: stats["2"] += 1
            elif wordConf >= 0.1: stats["1"] += 1
            else: stats["0"] += 1
            stats["parsedWords"] += 1

    # Initiate Document
    document = Document()
    # A4 Size
    document.sections[0].page_width = Mm(210)
    document.sections[0].page_height = Mm(297)
    # Font
    font = document.styles["Normal"].font
    font.name = "Calibri"

    # Document title and intro
    title = "Transcription Output"
    document.add_heading(title, level=1)

    # Intro
    document.add_paragraph("Transcription using AWS Transcribe automatic speech recognition, parsed by custom library.")
    document.add_paragraph()  # Spacing

    # Pull out header information just from the Transcribe job details
    table = document.add_table(rows=1, cols=2)
    table.style = document.styles["Light List Accent 1"]
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Job Name"
    hdr_cells[1].text = transcribeJjobInfo["TranscriptionJobName"]
    row_cells = table.add_row().cells
    row_cells[0].text = "Language"
    row_cells[1].text = analysisJobInfo["LanguageCode"]
    row_cells = table.add_row().cells
    row_cells[0].text = "File Format"
    row_cells[1].text = transcribeJjobInfo["MediaFormat"]
    row_cells = table.add_row().cells
    row_cells[0].text = "Sample Rate"
    row_cells[1].text = str(transcribeJjobInfo["MediaSampleRateHertz"]) + "Hz"
    row_cells = table.add_row().cells
    row_cells[0].text = "Audio Ident"
    if transcribeJjobInfo["ChannelIdentification"]:
        row_cells[1].text = "Channel-separated"
    else:
        row_cells[1].text = "Speaker-separated"
    row_cells = table.add_row().cells
    row_cells[0].text = "Vocabulary"
    if "VocabularyName" in transcribeJjobInfo:
        row_cells[1].text = transcribeJjobInfo["VocabularyName"]
    else:
        row_cells[1].text = "n/a"
    row_cells = table.add_row().cells
    row_cells[0].text = "Avg. Accuracy"
    row_cells[1].text = str(round(transcribeJjobInfo["AverageAccuracy"] * 100.0, 2)) + "%"
    row_cells = table.add_row().cells
    row_cells[0].text = "Parsed"
    row_cells[1].text = datetime.datetime.now().strftime("%A %d %B %Y at %X")
    document.add_paragraph()  # Spacing

    # Formatting transcript table widths
    widths = (Inches(1.2), Inches(3.0))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    # Display confidence count table - start with the fixed headers
    table = document.add_table(rows=1, cols=3)
    table.style = document.styles["Light List Accent 1"]
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Confidence"
    hdr_cells[1].text = "Count"
    hdr_cells[2].text = "Percentage"
    parsedWords = stats["parsedWords"]
    confidenceRanges = ["98% - 100%", "90% - 97%", "80% - 89%", "70% - 79%", "60% - 69%", "50% - 59%", "40% - 49%", "30% - 39%", "20% - 29%", "10% - 19%", "0% - 9%"]
    confidenceRangeStats = ["9.8", "9", "8", "7", "6", "5", "4", "3", "2", "1", "0"]

    # Add on each row
    for confRange, rangeStats in zip(confidenceRanges, confidenceRangeStats):
        addNextConfidenceRow(table, confRange, stats[rangeStats], parsedWords)

    # Add paragraph for spacing
    document.add_paragraph()

    # Confidence of each word as scatter graph
    plt.scatter(stats["timestamps"], stats["accuracy"])
    # Mean average as line across graph
    plt.plot([stats["timestamps"][0], stats["timestamps"][-1]], [statistics.mean(stats["accuracy"]), statistics.mean(stats["accuracy"])], "r")
    # Formatting
    plt.xlabel("Time (seconds)")
    # plt.xticks(range(0, int(stats['timestamps'][-1]), 60))
    plt.ylabel("Accuracy (percent)")
    plt.yticks(range(0, 101, 10))
    plt.title("Accuracy During Transcription")
    plt.legend(["Accuracy average (mean)", "Individual words"], loc="lower center")

    # Write out the chart
    accuracy_chart_file_name = "./" + "chart.png"
    plt.savefig(accuracy_chart_file_name)
    plt.clf()
    document.add_picture(accuracy_chart_file_name, width=Cm(14.64))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()

    # Process and display transcript by speaker segments
    table = document.add_table(rows=1, cols=6)
    table.style = document.styles["Light List Accent 1"]
    hdr_cells = table.rows[0].cells
    hdr_cells[COL_STARTTIME].text = "Start"
    hdr_cells[COL_ENDTIME].text = "End"
    hdr_cells[COL_SPEAKER].text = "Speaker"
    hdr_cells[COL_CONTENT].text = "Content"
    hdr_cells[COL_SENTIMENT].text = "Sentiment"
    hdr_cells[COL_SENTIMENT_SCORE].text = "Score"

    # Based upon our segment list, write out the transcription table
    writeOutTranscribeTable(table, speechSegmentList)

    # Formatting transcript table widthds
    widths = (Inches(0.8), Inches(0.8), Inches(0.8), Inches(4.5), Inches(0.8), Inches(0.4))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    # Generate sentiment graphs, starting by pulling out our two data streams for just pos/neg items
    speaker0labels = ['ch_0', 'spk_0']
    speaker1labels = ['ch_1', 'spk_1']
    speaker0timestamps = []
    speaker0data = []
    speaker1timestamps = []
    speaker1data = []

    # Generate our raw data
    for segment in speechSegmentList:
        if bool(segment["SentimentIsPositive"]) or bool(segment["SentimentIsNegative"]):
            # Only interested in actual sentiment entries
            timestamp = float(segment["SegmentStartTime"])

            # Positive re-calculation
            if bool(segment["SentimentIsPositive"]):
                score = 2 * ((1-(1-float(segment["SentimentScore"]))/(1 - MIN_SENTIMENT_POSITIVE))*0.5)
            # Negative re-calculation
            else:
                score = 2 * ((1-float(segment["SentimentScore"]))/(1 - MIN_SENTIMENT_NEGATIVE)*0.5-0.5)

            if segment["SegmentSpeaker"] in speaker1labels:
                speaker1data.append(score)
                speaker1timestamps.append(timestamp)
            elif segment["SegmentSpeaker"] in speaker0labels:
                speaker0data.append(score)
                speaker0timestamps.append(timestamp)
            else:
                # DEBUG - shouldn't happen
                print("Couldn't find " + segment.segmentSpeaker)

    # Spline fit needs at least 4 points for k=3, but 5 works better
    speaker1k = 3
    speaker0k = 3
    if len(speaker1data) < 5:
        speaker1k = 1
    if len(speaker0data) < 5:
        speaker0k = 1

    # Creater Speaker-0 graph
    plt.figure(figsize=(8, 5))
    speaker0xnew = np.linspace(speaker0timestamps[0], speaker0timestamps[-1], int((speaker0timestamps[-1] - speaker0timestamps[0]) + 1.0))
    speaker0spl = make_interp_spline(speaker0timestamps, speaker0data, k=speaker0k)
    speaker0powerSmooth = speaker0spl(speaker0xnew)
    plt.plot(speaker0timestamps, speaker0data, "ro")
    plt.plot(speaker0xnew, speaker0powerSmooth, "r", label="Speaker 1")

    # Create Speaker-1 graph
    speaker1xnew = np.linspace(speaker1timestamps[0], speaker1timestamps[-1], int((speaker1timestamps[-1] - speaker1timestamps[0]) + 1.0))
    speaker1spl = make_interp_spline(speaker1timestamps, speaker1data, k=speaker1k)
    speaker1powerSmooth = speaker1spl(speaker1xnew)
    plt.plot(speaker1timestamps, speaker1data, "bo")
    plt.plot(speaker1xnew, speaker1powerSmooth, "b", label="Speaker 2")

    # Draw it out
    plt.title("Call Sentiment - Pos/Neg Only")
    plt.xlabel("Time (seconds)")
    plt.axis([0, max(speaker0timestamps[-1], speaker1timestamps[-1]), -1.5, 1.5])
    plt.legend()
    plt.axhline(y=0, color='k')
    plt.axvline(x=0, color='k')
    plt.grid(True)
    plt.xticks(np.arange(0, max(speaker0timestamps[-1], speaker1timestamps[-1]), 60))
    plt.yticks(np.arange(-1, 1.01, 0.25))

    # Write out the chart
    sentiment_chart_file_name = "./" + "sentiment.png"
    plt.savefig(sentiment_chart_file_name)
    plt.clf()
    document.add_picture(sentiment_chart_file_name, width=Cm(14.64))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save the whole document
    document.save(docxFilename)

    # Now remove the helper image files, which we don't need
    os.remove(accuracy_chart_file_name)
    os.remove(sentiment_chart_file_name)

def createDevReloadList():
    return [
        "0a.93.a0.3e.00.00-09.26.37.755-09-23-2019.wav",
        "0a.93.a0.3e.00.00-09.11.32.483-09-10-2019.wav",
        "0a.93.a0.3e.00.00-09.28.29.553-09-17-2019.wav",
        "0a.93.a0.3e.00.00-09.28.52.023-09-10-2019.wav",
        "0a.93.a0.3f.00.00-10.41.54.226-09-20-2019.wav",
        "0a.93.a0.3f.00.00-10.46.53.432-09-19-2019.wav",
        "0a.93.a0.3e.00.00-09.31.33.923-09-16-2019.wav",
        "0a.93.a0.3e.00.00-09.30.26.530-09-05-2019.wav",
        "0a.93.a0.3e.00.00-09.25.51.067-09-26-2019.wav",
        "0a.93.a0.3e.00.00-09.13.43.164-09-16-2019.wav"
    ]


def generateDocument():

    # Parameter extraction
    assert len(sys.argv) > 1, "Usage: tswrite {transcribeJobName} [--reload]]"
    transcribeJobName = sys.argv[1]
    fullReload = False
    if len(sys.argv) > 2:
        if sys.argv[2] == "--reload":
            fullReload = True

    # Workflow for process
    start = perf_counter()

    if not fullReload:
        # Just parsing a single a single document, so create a Word Doc on from the results
        transcribeParser = TranscribeParser(MIN_SENTIMENT_POSITIVE, MIN_SENTIMENT_NEGATIVE, COMPREHEND_ENTITY)
        transcribeParser.parseTranscribeFile(transcribeJobName)

        # Now write the JSON results to a local file, and create the Word document from that
        jsonOutputFilename = transcribeParser.getJSONOutputFilename()
        if len(sys.argv) > 2:
            outputDocName = sys.argv[2]
        else:
            outputDocName = jsonOutputFilename + ".docx"
        with open(jsonOutputFilename, 'w') as fileWrite:
            json.dump(transcribeParser.outputAsJSON(), fileWrite)
        write(jsonOutputFilename, outputDocName, transcribeParser)
        print(f"> Transcript {outputDocName} writen.")
    else:
        # Full reload - no Word output, just pile through them
        reloadList = createDevReloadList()
        print(f"Bulk re-processing: {len(reloadList)} files to process")
        for transcribeJobName in reloadList:
            print(f"Processing {transcribeJobName}...")
            transcribeParser = TranscribeParser(MIN_SENTIMENT_POSITIVE, MIN_SENTIMENT_NEGATIVE, COMPREHEND_ENTITY)
            transcribeParser.parseTranscribeFile(transcribeJobName)

    # Write out the performance statistics
    finish = perf_counter()
    duration = round(finish - start, 2)
    print(f"> Processing complete in {duration} seconds.")


# Main entrypoint
if __name__ == "__main__":
    generateDocument()
